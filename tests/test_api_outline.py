"""Tests for the outline API endpoint (LLM is mocked)."""
import json

import pytest
from docx import Document
from fastapi.testclient import TestClient
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, NameObject

from app.config_store import ConfigStore
from app.db import init_db
from app.main import app
from app.models import ModelConfig


@pytest.fixture
def client_with_outline_config(tmp_path, monkeypatch):
    db = tmp_path / "app.db"
    init_db(db)
    monkeypatch.setenv("PPTM_DATA_DIR", str(tmp_path))
    monkeypatch.setenv("PPTM_OUTPUTS_DIR", str(tmp_path / "outputs"))
    (tmp_path / "outputs").mkdir()
    ConfigStore(db_path=db).save(
        ModelConfig(stage="outline", provider="openai", api_key="k", model_name="gpt-test")
    )
    with TestClient(app) as c:
        yield c


def _stub_provider(monkeypatch):
    from app import main

    class StubProvider:
        def chat(self, system, user, *, temperature=0.7, max_tokens=2048):
            return json.dumps({
                "pages": [{"title": "P1", "key_points": ["a"], "layout": "title-content"}]
            })

    monkeypatch.setattr(main, "build_provider", lambda cfg: StubProvider())


def test_generate_outline_requires_config(tmp_path, monkeypatch):
    monkeypatch.setenv("PPTM_DATA_DIR", str(tmp_path))
    monkeypatch.setenv("PPTM_OUTPUTS_DIR", str(tmp_path / "outputs"))
    (tmp_path / "outputs").mkdir()
    init_db(tmp_path / "test.db")
    with TestClient(app) as c:
        r = c.post("/api/outline/generate", data={"topic": "x"})
    assert r.status_code == 400


def test_generate_outline_returns_id_and_content(client_with_outline_config, monkeypatch):
    _stub_provider(monkeypatch)
    r = client_with_outline_config.post("/api/outline/generate", data={"topic": "AI"})
    assert r.status_code == 200
    data = r.json()
    assert data["outline_id"] >= 1
    assert data["content"]["topic"] == "AI"
    assert data["content"]["pages"][0]["title"] == "P1"
    assert data["source_files"] == []


def test_generate_outline_rejects_empty_topic(client_with_outline_config, monkeypatch):
    _stub_provider(monkeypatch)
    r = client_with_outline_config.post("/api/outline/generate", data={"topic": "   "})
    assert r.status_code == 400


def test_generate_outline_with_docx_attachment(client_with_outline_config, monkeypatch, tmp_path):
    _stub_provider(monkeypatch)
    from app import main as app_main
    captured = {}

    class CapturingProvider:
        def chat(self, system, user, *, temperature=0.7, max_tokens=2048):
            captured["user"] = user
            return json.dumps({
                "pages": [{"title": "P1", "key_points": ["a"], "layout": "title-content"}]
            })

    monkeypatch.setattr(app_main, "build_provider", lambda cfg: CapturingProvider())

    docx_path = tmp_path / "source.docx"
    doc = Document()
    doc.add_paragraph("Quarterly revenue grew 30%.")
    doc.add_paragraph("Main product lines are A and B.")
    docx_path.write_bytes(b"")  # placeholder; we'll save below
    doc.save(str(docx_path))

    with docx_path.open("rb") as f:
        r = client_with_outline_config.post(
            "/api/outline/generate",
            data={"topic": "Q4 report"},
            files={"files": ("source.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
    assert r.status_code == 200
    data = r.json()
    assert len(data["source_files"]) == 1
    sf = data["source_files"][0]
    assert sf["filename"] == "source.docx"
    assert sf["char_count"] > 0
    assert sf["stored_path"].endswith("source.docx")
    # Verify document text was included in the LLM prompt
    assert "Quarterly revenue" in captured["user"]
    assert "=== source.docx ===" in captured["user"]


def test_generate_outline_with_unsupported_file_400(client_with_outline_config, tmp_path):
    txt_path = tmp_path / "notes.txt"
    txt_path.write_text("hello", encoding="utf-8")
    with txt_path.open("rb") as f:
        r = client_with_outline_config.post(
            "/api/outline/generate",
            data={"topic": "T"},
            files={"files": ("notes.txt", f, "text/plain")},
        )
    assert r.status_code == 400
    assert "Unsupported" in r.json()["detail"]


def test_generate_outline_with_multiple_files(client_with_outline_config, monkeypatch, tmp_path):
    _stub_provider(monkeypatch)
    from app import main as app_main
    captured = {}

    class CapturingProvider:
        def chat(self, system, user, *, temperature=0.7, max_tokens=2048):
            captured["user"] = user
            return json.dumps({
                "pages": [{"title": "P1", "key_points": [], "layout": "title-content"}]
            })

    monkeypatch.setattr(app_main, "build_provider", lambda cfg: CapturingProvider())

    docx_path = tmp_path / "a.docx"
    doc_a = Document()
    doc_a.add_paragraph("From A")
    doc_a.save(str(docx_path))
    pdf_path = tmp_path / "b.pdf"
    w = PdfWriter()
    page = w.add_blank_page(width=612, height=792)
    s = DecodedStreamObject()
    s.set_data(b"BT /F1 12 Tf 50 750 Td (From B) Tj ET")
    page[NameObject("/Contents")] = w._add_object(s)
    with pdf_path.open("wb") as f:
        w.write(f)

    with docx_path.open("rb") as fa, pdf_path.open("rb") as fb:
        r = client_with_outline_config.post(
            "/api/outline/generate",
            data={"topic": "Combined"},
            files=[
                ("files", ("a.docx", fa, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")),
                ("files", ("b.pdf", fb, "application/pdf")),
            ],
        )
    assert r.status_code == 200
    data = r.json()
    assert len(data["source_files"]) == 2
    assert {sf["filename"] for sf in data["source_files"]} == {"a.docx", "b.pdf"}
    assert "From A" in captured["user"]
    # The synthetic PDF has no extractable text (pypdf needs a font dict for Tj),
    # so we only verify it was attached; text extraction is exercised in the
    # dedicated parser tests.
    assert "=== b.pdf ===" in captured["user"]


def test_history_outlines_includes_source_files(client_with_outline_config, monkeypatch, tmp_path):
    _stub_provider(monkeypatch)
    docx_path = tmp_path / "h.docx"
    doc_h = Document()
    doc_h.add_paragraph("History content")
    doc_h.save(str(docx_path))
    with docx_path.open("rb") as f:
        r = client_with_outline_config.post(
            "/api/outline/generate",
            data={"topic": "H"},
            files={"files": ("h.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
    assert r.status_code == 200
    oid = r.json()["outline_id"]

    hist = client_with_outline_config.get("/api/history/outlines").json()
    assert hist[0]["id"] == oid
    assert hist[0]["source_files"][0]["filename"] == "h.docx"


def test_download_source_file(client_with_outline_config, monkeypatch, tmp_path):
    _stub_provider(monkeypatch)
    docx_path = tmp_path / "dl.docx"
    doc_dl = Document()
    doc_dl.add_paragraph("Downloadable content")
    doc_dl.save(str(docx_path))
    with docx_path.open("rb") as f:
        r = client_with_outline_config.post(
            "/api/outline/generate",
            data={"topic": "D"},
            files={"files": ("dl.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
    oid = r.json()["outline_id"]

    r = client_with_outline_config.get(f"/api/outline/{oid}/source/dl.docx")
    assert r.status_code == 200
    assert r.headers["content-type"].startswith("application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    assert len(r.content) > 0


def test_download_source_file_404_for_missing(client_with_outline_config, monkeypatch, tmp_path):
    _stub_provider(monkeypatch)
    r = client_with_outline_config.post("/api/outline/generate", data={"topic": "x"})
    oid = r.json()["outline_id"]
    r = client_with_outline_config.get(f"/api/outline/{oid}/source/missing.pdf")
    assert r.status_code == 404